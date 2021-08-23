import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
import sys,os
from tkcalendar import Calendar, DateEntry
import random,string
from datetime import date
from tkinter.filedialog import askdirectory
import docx
import docx2pdf
import webbrowser
from PIL import ImageTk,Image 

win= tk.Tk()

_bgcolor = '#d9d9d9'  # X11 color: 'gray85'
_fgcolor = '#000000'  # X11 color: 'black'
_compcolor = '#d9d9d9' # X11 color: 'gray85'
_ana1color = '#d9d9d9' # X11 color: 'gray85'
_ana2color = '#ececec' # Closest X11 color: 'gray92'
######### All Images ########
Register_New_Photo = ImageTk.PhotoImage(Image.open(r"All_image\Register_customer.jpg"))
Home_image = ImageTk.PhotoImage(Image.open(r"All_image\base12.jpg"))
New_After_login_Photo = ImageTk.PhotoImage(Image.open(r"All_image\after_login.jpg"))
Book_trip_button = ImageTk.PhotoImage(Image.open(r'All_image\History_button.jpg'))
History_button = ImageTk.PhotoImage(Image.open(r'All_image\history_b.jpg'))
apply_for_job_photo=ImageTk.PhotoImage(Image.open(r"All_image\apply_for_job_photo.jpg"))
Drive_register= ImageTk.PhotoImage(Image.open(r"All_image\driver_register.jpg"))
Book_page = ImageTk.PhotoImage(Image.open(r"All_image\booking_page1.jpg"))
admin_login = ImageTk.PhotoImage(Image.open(r"All_image\admin_login.jpg"))
Admin_menu_ = ImageTk.PhotoImage(Image.open(r"All_image\admin_menu.jpg"))
Display_driver = ImageTk.PhotoImage(Image.open(r"All_image\display_driver.jpg"))
Display_Users = ImageTk.PhotoImage(Image.open(r"All_image\display_Users.jpg"))
Display_history = ImageTk.PhotoImage(Image.open(r"All_image\display_history.jpg"))
Add_driver= ImageTk.PhotoImage(Image.open(r"All_image\add_New_driver.jpg"))
edit = ImageTk.PhotoImage(Image.open(r"All_image\edit.jpg"))
Display_history = ImageTk.PhotoImage(Image.open(r"All_image\display_history.jpg"))


win.geometry("1920x1017+660+210")
win.minsize(1924, 1061)
win.maxsize(1924, 1061)
win.resizable(1,  1)
win.title("Book My Trip")
win.configure(background="#d9d9d9")
win.wm_iconbitmap(r"taxi.ico")

Frame2 = tk.Frame(win)
Frame2.place(relx=0.0, rely=0.0, relheight=1.062, relwidth=1.0)
Frame2.configure(relief='groove')
Frame2.configure(borderwidth="2")
Frame2.configure(relief="groove")
Frame2.configure(background="#d9d9d9")
####################### HOME PAGE ####################
def New_home():
    for widget in Frame2.winfo_children():
            widget.destroy()

    Label1 = tk.Label(Frame2)
    Label1.place(relx=0.0, rely=0.0, height=1080, width=1920)
    Label1.configure(background="#000000")
    Label1.configure(disabledforeground="#a3a3a3")
    Label1.configure(foreground="#ffffff")
    Label1.configure(image=Home_image)
    Label1.configure(text='''Label''')


    u_email=tk.StringVar(value='Email')
    _uemail = tk.Entry(Frame2,textvariable=u_email)                          #Entry for user email
    _uemail.place(x=58,y=550,height=50,width=525)
    _uemail.configure(takefocus="",cursor="ibeam",relief='flat',font=('Helvetica', 32,'bold'),background="white",foreground="black") 
    _uemail.configure(takefocus="",cursor="ibeam #000000",insertbackground='Black')

    u_password=tk.StringVar()
    _upass = tk.Entry(Frame2,textvariable=u_password,show='*')              #Entry for user Password
    _upass.place(x=58,y=666,height=50,width=523)
    _upass.configure(takefocus="",cursor="ibeam",font=('Helvetica', 40,'bold'),background="white",foreground="black")
    _upass.configure(takefocus="",cursor="ibeam #000000",insertbackground='Black')

    def Log_in_satus():                                                     # Check User email and password In File
        with open("user_data.txt",'r') as ud:
            data=ud.read()
            data=data.split('\n')
            count1=1
            
            for line in data:
                # print(line)
                # print(len(line))
                line=tuple(line.split(' '))
                # print(line)
                # print(len(line))
                if count1<=len(data) and len(line[0])>1 and line[0]!='/n' and line[0]!=' ':
                    if u_email.get() == line[5] :
                        if u_password.get() == line[1]:
                            print("you are in")
                            New_After_login(u_email.get(),u_password.get())
                            return None
                        else:
                            tk.messagebox.showerror("Wrong Entry", "Check user name and password again")        # Warning 
                        count1+=1
                        continue                    
                    else:
                        if u_email.get() == line[0]:
                            tk.messagebox.showerror("Wrong Entry", "Enter User Email\n Not User name ")
                            return None
                    count1 +=1
                    continue
                elif count1 >= len(data):
                    tk.messagebox.showerror("Wrong Entry", "Check user name and password again")
                    return None
                count1+=1
            
            
    TButton1 = tk.Button(Frame2, command=Log_in_satus)
    TButton1.place(x=380,y=804, height=50, width=178)
    TButton1.configure(takefocus="",text='''Submit''',font=('Copperplate Gothic Bold', 28,'bold'),background="White",relief='flat',foreground="#7ea8ab")

    Signup = tk.Button(Frame2, command=lambda: Register_New())
    Signup.place(x=92,y=804, height=50, width=178)
    Signup.configure(takefocus="",text='''Sign Up''',font=('Copperplate Gothic Bold', 28,'bold'),background="White",relief='flat',foreground="#7ea8ab")

    Signup = tk.Button(Frame2, command=lambda: admin_log_in())
    Signup.place(x=1,y=960, height=25, width=35)
    Signup.configure(takefocus="",text='''?''',font=('Copperplate Gothic Bold', 28,'bold'),background="#323436",relief='flat',foreground="#7ea8ab")
####################### After Log in #################
def New_After_login(uemail,password):                                   #After LOGIN main menu for client
        for widget in Frame2.winfo_children():
                widget.destroy()

        Label1 = tk.Label(Frame2)
        Label1.configure(background="#000000")
        Label1.configure(disabledforeground="#a3a3a3")
        Label1.configure(foreground="#ffffff")
        Label1.configure(text='''Label''')
        Label1.place(relx=0.0, rely=0.0, height=1080, width=1920)
        Label1.configure(image=New_After_login_Photo)

        # resized = original.resize((468, 65),Image.ANTIALIAS)
        # image = ImageTk.PhotoImage(resized)
        bookin = tk.Button(Frame2,command=lambda: book_trip(uemail,password))
        bookin.place(x=727,y=511, height=56, width=467)
        bookin.configure(image=Book_trip_button,relief="flat",border=0)

        history = tk.Button(Frame2,command=lambda: History_user(uemail,password))
        history.place(x=769,y=602, height=59, width=380)
        history.configure(image=History_button,relief="flat",border=0)

        back = tk.Button(Frame2, command=lambda: New_home())
        back.place(x=865,y=695, height=65, width=194)
        back.configure(takefocus="",text='''BACK''',font=('Copperplate Gothic Bold', 34,'bold'),background="white",relief='flat',foreground="#7ea8ab")

        Apply_for_job = tk.Button(Frame2, command=lambda: New_Driver(uemail,password))
        Apply_for_job.place(x=1052,y=846, height=20, width=185)
        Apply_for_job.configure(image=apply_for_job_photo,relief='flat')
####################### Book Texi ####################
def book_trip(uemail,password):                                                 # Booking of Trip and genrating bill and saving history
    for widget in Frame2.winfo_children():
        widget.destroy()
    Label1 = tk.Label(Frame2)
    Label1.place(relx=0.0, rely=0.0, height=1080, width=1920)
    Label1.configure(background="#6f6f64")
    Label1.configure(disabledforeground="#a3a3a3")
    Label1.configure(foreground="#ffffff")
    Label1.configure(image=Book_page)
    Label1.configure(text='''Label''')

    uname=''

    # with open("user_data.txt",'r') as ud:
    #         data=ud.read()
    #         data=data.split('\n')
    #         count1=1
    #         for line in data:
    #             line=tuple(line.split(' '))
    #             if count1<=len(data):
    #                 if uemail == line[5] :
    #                     uname=line[0]
    #                 count1 +=1

    with open("user_data.txt",'r') as ud:
        for i in ud.readlines():
            if len(i) > 50 and i != '\n':
                    i=i.split(' ')
                    if uemail==i[5]:
                        uname=i[0]
                        break 

    V_Name=tk.StringVar(value=uname)
    E_Name = tk.Entry(Frame2,textvariable=V_Name,state=tk.DISABLED)
    E_Name.place(x=490,y=117, height=48, width=247)
    E_Name.configure(disabledforeground="#6f6f64",disabledbackground="#cbd9ca",relief='flat',foreground="#000000",insertbackground="#6f6f64",font=('Helvetica', 32,'bold')) #Maulik_GUI

    city_source=[]
    city_destination=[]
    with open("places.txt",'r') as ud:
        data=ud.read()
        data=data.split('\n')
        for i in data:
            i=i.split(' ')
            if i[0]!=' ':
                city_source.append(i[0])
                city_destination.append(i[0])
        city_source=tuple(set(city_source))
        city_destination=tuple(set(city_destination))
        # print(city_destination)
        # print(city_source)
    def change_source(object):
        p=S_Combo.get()
        s.configure(text=p)

    s_index=tk.StringVar()
    S_Combo = ttk.Combobox(Frame2,textvariable=s_index)
    S_Combo.place(x=490,y=198, height=48, width=247)
    S_Combo['values']=city_source
    S_Combo.current(2)
    S_Combo.bind("<<ComboboxSelected>>",change_source)

    p=S_Combo.get()
    s=ttk.Label(Frame2)
    s.place(x=490,y=198,height=48,width=230)
    s.configure(text=p,font=('Helvetica', 28,'bold'),relief='flat',foreground="#6f6f64",background="#cbd9ca")

    def change_destination(object):
        desti=D_Combo.get()
        d.configure(text=desti)


    d_index=tk.StringVar()
    D_Combo = ttk.Combobox(Frame2,textvariable=d_index)
    D_Combo.place(x=490, y=279,height=48, width=247)
    D_Combo['values']=city_destination
    D_Combo.current(0)
    D_Combo.bind("<<ComboboxSelected>>",change_destination)

    desti=D_Combo.get()
    d=ttk.Label(Frame2)
    d.place(x=490,y=279,height=48,width=230)
    d.configure(text=desti,font=('Helvetica', 28,'bold'),relief='flat',foreground="#6f6f64",background="#cbd9ca")

    def show_date(object):
        w=str(booking_date.get())
        w=w.split('/')
        p=w[1]+'/'+w[0]+'/'+w[2]
        hide_by_lable1.configure(text=p,font=('Helvetica', 20,'bold'))
    today=date.today()
    booking_date=tk.StringVar()
    cal = DateEntry(Frame2,textvariable=booking_date,width=12, background='darkblue',foreground='white', borderwidth=2, month=today.month,year=today.year,day=today.day)
    cal.place(x=490, y=360,height=48, width=247)
    cal.bind("<<DateEntrySelected>>",show_date)


    w=str(today.day)+"/"+str(today.month)+"/"+str(today.year)
    hide_by_lable1=tk.Label(Frame2)
    hide_by_lable1.place(x=490, y=360,height=48, width=230)
    hide_by_lable1.configure(text=w,background='#cbd9ca',foreground='#6f6f64',font=('Helvetica', 20,'bold'),relief='flat')


    time1=tk.StringVar()
    Entry1_1 = tk.Entry(Frame2,textvariable=time1)
    Entry1_1.place(x=490, y=441,height=48, width=247)
    Entry1_1.configure(background="#cbd9ca",foreground="#6f6f64",insertbackground="#6f6f64",font=('Helvetica', 20,'bold'),relief='flat')


    type_booking=tk.IntVar()
    Radiobutton1 = tk.Radiobutton(Frame2,command=lambda:refresh())
    Radiobutton1.place(x=540, y=505,height=28, width=140)
    Radiobutton1.configure(activebackground="#cbd9ca",activeforeground="#6f6f64",background='#cbd9ca',foreground='#6f6f64',font=('Helvetica', 20,'bold'),relief='flat',text='''Regular''')
    Radiobutton1.configure(variable=type_booking,value=0)

    Radiobutton1_1 = tk.Radiobutton(Frame2,command=lambda:refresh())
    Radiobutton1_1.place(x=540, y=540,height=28, width=175)
    Radiobutton1_1.configure(activebackground="#cbd9ca",activeforeground="#6f6f64",background='#cbd9ca',foreground='#6f6f64',font=('Helvetica', 20,'bold'),relief='flat',text='''Luxurious''')
    Radiobutton1_1.configure(variable=type_booking,value=1)

    check=tk.IntVar()
    Radiobutton1_2 = tk.Checkbutton(Frame2,command=lambda:refresh())
    Radiobutton1_2.place(x=540, y=575,height=28, width=130)
    Radiobutton1_2.configure(activebackground="#cbd9ca",activeforeground="#6f6f64",background='#cbd9ca',foreground='#6f6f64',font=('Helvetica', 20,'bold'),relief='flat',text='''Return''')
    Radiobutton1_2.configure( onvalue=1, offvalue=0,variable=check)


    get_distance = tk.Label(Frame2)
    get_distance.place(x=490, y=622, height=48, width=247)
    get_distance.configure(background="#cbd9ca",foreground="#6f6f64",font=('Helvetica', 20,'bold'),relief='flat')

    get_price = tk.Label(Frame2)
    get_price.place(x=490, y=734, height=48, width=247)
    get_price.configure(background="#cbd9ca",foreground="#6f6f64",font=('Helvetica', 20,'bold'),relief='flat')


    def refresh():                                          # fatch data of price and Distance
        with open("places.txt",'r') as ud:
            source=s_index.get()
            destination=d_index.get()
            data=ud.read()
            data=data.split('\n')
            count1=1
            dist=" "
            price=' '
            for line in data:
                line=tuple(line.split(' '))
                if count1<=len(data):
                    if source in line and destination in line:
                        dist=f"{line[2]} Km"
                        if type_booking.get()==0:
                            price=f"\u20B9 {line[3]}"
                        if type_booking.get()==1:
                            price=f"\u20B9 {line[4]}"
                        if check.get()==1:
                            if type_booking.get()==0:
                                price=f"\u20B9 {int(line[3])*2-50}"
                            if type_booking.get()==1:
                                price=f"\u20B9 {int(line[4])*2-100}"
                        break
                    count1+=1
            get_distance.config(text=dist)
            get_price.config(text=price)
    # B_refersh = tk.Button(Frame2,command= refresh)
    # B_refersh.place(relx=0.728, rely=0.323, height=24, width=47)
    # B_refersh.configure(activebackground="#ececec",activeforeground="#6f6f64",background="#d9d9d9",disabledforeground="#a3a3a3",foreground="#6f6f64",highlightbackground="#d9d9d9",highlightcolor="#6f6f64",pady="0",text='''Refresh''')

    def generat_bill2(name,source,destination,Date,Time,type_booking,Check=0,price1=0,dist=0,driver='None'):        #Generation of Bill in pdf format
        doc=docx.Document("Base.docx")
        table =doc.add_table(rows=1,cols=2)
        with open("places.txt",'r', encoding='utf-8') as ud:
            data=ud.read()
            data=data.split('\n')
            count1=1
            for line in data:
                line=tuple(line.split(' '))
                if count1<=len(data):
                    if source in line and destination in line:
                        dist=f"{line[2]} Km"
                        if type_booking==0:
                            price1=f" \u20B9 {line[3]}"
                        if type_booking==1:
                            price1=f" \u20B9 {line[4]}"
                        if Check==1:
                            if type_booking==0:
                                price1=f" \u20B9 {int(line[3])*2-50}"
                            if type_booking==1:
                                price1=f" \u20B9 {int(line[4])*2-100}"
                        break
                    count1+=1
            if type_booking==0:
                type_booking1="Regular"
            if type_booking==1:
                type_booking1="Luxurious"
            if Check==1:
                if type_booking==0:
                    type_booking1="Regular" +" Return"
                if type_booking==1:
                    type_booking1="Luxurious"+" Return"
            
        with open("driver.txt",'r') as f:                                   # Assigning Driver 
            d_name=f.read()
            d_name=d_name.split('\n')
            i=random.randint(0,len(d_name)-2)
            d_name=d_name[i].split(' ')
            driver =d_name[0]
        
        car_number="GJ "+str(random.randint(1,36))+" "+ str(random.choice(string.ascii_letters)).upper() +' ' +str(random.randint(1000,9999))
        ref=str(random.randint(10000,99999))
        data={
            "Reference NO" :ref,
            "Name" : name,
            "Source" : source,
            "Destination" : destination,
            "Date & Time" : f"{Date} & {Time}",
            "Type" : str(type_booking1),
            "Price" : str(price1),
            "Total Distance" : dist,
            "Drive Name" : driver,
            "Car Number" : car_number
        }
        # row = table.rows[1]
        # row.cells[0].text = 'Foo bar to you.'
        # row.cells[1].text = 'And a hearty foo bar to you too sir!'
        for field,value in data.items():
            row_cells=table.add_row().cells
            row_cells[0].text=str(field)
            row_cells[1].text=str(value)
        paragraph = doc.add_paragraph('')
        values1= paragraph.add_run("\n\t\tPayMent:: Case after Journy")
        values1.font.name = 'Helvetica'
        values1.font.size = docx.shared.Pt(20)
        values= paragraph.add_run("\n\t\t\tThank you\n \t\tHave a great Journy")
        values.font.name = 'Helvetica'
        values.font.size = docx.shared.Pt(35)
        filename=f"{data['Name']} {ref}"
        doc.save(f"{filename}.docx")
        with open("history.txt",'a', encoding='utf-8') as hi:
            for i in data.keys():
                hi.write(data[i]+',')
            hi.write("\n")
        ####### save file#####
        url=askdirectory()
        docx2pdf.convert(f"{filename}.docx",f"{url}\\{filename}.pdf")
        os.remove(f"{filename}.docx")

    B_generat_bill = tk.Button(Frame2,command=lambda: generat_bill2(V_Name.get(),s_index.get(),d_index.get(),booking_date.get(),time1.get(),type_booking.get(),check.get()))
    B_generat_bill.place(x=491,y=838, height=47, width=247)
    B_generat_bill.configure(activebackground="#cbd9ca",activeforeground="#6f6f64",background="#cbd9ca",disabledforeground="#a3a3a3",font=('Helvetica', 28,'bold'),foreground="#6f6f64",highlightbackground="#d9d9d9",highlightcolor="#6f6f64",relief='flat',text='''Generat Bill''')

    B_back = tk.Button(Frame2,command=lambda: New_After_login(uemail,password))
    B_back.place(x=196,y=838, height=47, width=247)
    B_back.configure(activebackground="#cbd9ca",activeforeground="#6f6f64",background="#cbd9ca",disabledforeground="#a3a3a3",font=('Helvetica', 28,'bold'),foreground="#6f6f64",highlightbackground="#d9d9d9",highlightcolor="#6f6f64",relief='flat',text='''Back''')

    def goto():
        base_url='https://www.google.com/maps/dir/'
        url=f'{base_url}{s_index.get()}/{d_index.get()}'
        if S_Combo.get():
            webbrowser.open(url, new=2)


    Button1 = tk.Button(Frame2,command= goto)
    Button1.place(x=315,y=926, height=47, width=305)
    Button1.configure(activebackground="#cbd9ca",activeforeground="#6f6f64",background="#cbd9ca",disabledforeground="#a3a3a3",font=('Helvetica', 25,'bold'),foreground="#6f6f64",highlightbackground="#d9d9d9",highlightcolor="#6f6f64",relief='flat',text='''See Route On Map''')
####################### New Customer #################
def Register_New():                                                             #Client Sign Up forum
    for widget in Frame2.winfo_children():
        widget.destroy()
    Label1 = tk.Label(Frame2)
    Label1.place(relx=0.0, rely=0.0, height=1080, width=1920)
    Label1.configure(background="#d9d9d9")
    Label1.configure(disabledforeground="#a3a3a3")
    Label1.configure(foreground="#000000")

    Label1.configure(image=Register_New_Photo)
    Label1.configure(text='''Label''')

    ########Entry Boxes########
    name=tk.StringVar()
    Entry1_1 = tk.Entry(Frame2,textvariable=name)
    Entry1_1.place(x=447, y=141, height=34, width=342)
    Entry1_1.configure(background="#ffffff")
    Entry1_1.configure(font=('Helvetica', 20))
    Entry1_1.configure(foreground="#000000")
    Entry1_1.configure(relief="flat")


    password=tk.StringVar()
    password_entry = tk.Entry(Frame2, show="*",textvariable=password)
    password_entry.place(x=447, y=227, height=34, width=342)
    password_entry.configure(background="#ffffff",show="*")
    password_entry.configure(font=('Helvetica', 20))
    password_entry.configure(foreground="#000000")
    password_entry.configure(relief="flat")
    password_entry.configure(takefocus="",cursor="ibeam")

    age=tk.IntVar()
    Spinbox1 = tk.Spinbox(Frame2, from_=1.0, to=120.0)
    Spinbox1.place(x=447, y=313, height=34, width=342)
    Spinbox1.configure(activebackground="#f9f9f9",buttonbackground="#d9d9d9",background="white",disabledforeground="#a3a3a3",font=('Helvetica', 20),relief='flat',foreground="black",highlightbackground="black",highlightcolor="black",insertbackground="black",selectbackground="#00ffff",selectforeground="white")
    Spinbox1.configure(textvariable=age)



    gender=tk.StringVar()
    TCombobox1 = ttk.Combobox(Frame2,style="TCombobox")
    TCombobox1.place(x=447, y=400, height=33,width=342)
    TCombobox1.configure(textvariable=gender,font=('Helvetica', 20))
    TCombobox1['values']=('Male','Female','Other')
    TCombobox1.current(0)
    ttk.Style().configure('TCombobox',foregroundcolor='black', relief='flat',borderwidth=10)
    # lable4=tk.Label(Frame2)
    # lable4.place(x=470,y=403,height=30,width=100)
    # lable4.configure(image=transperent,relief="flat")
    # male=tk.Radiobutton(Frame2)
    # male.place(x=444, y=403)
    # male.configure(textvariable=gender)

    # female=tk.Radiobutton(Frame2)
    # female.place(x=580, y=403)
    # female.configure(textvariable=gender)


    hide1=tk.Label(Frame2)
    hide1.place(x=447,y=400,width=342,height=1)
    hide1.configure(background="white")

    hide2=tk.Label(Frame2)
    hide2.place(x=447,y=398,width=2,height=34)
    hide2.configure(background="white")

    hide3=tk.Label(Frame2)
    hide3.place(x=447,y=432,width=342,height=2)
    hide3.configure(background="white")

    hide4=tk.Label(Frame2)
    hide4.place(x=788,y=400,width=5,height=33)
    hide4.configure(background="white")

    phon_no=tk.IntVar()
    TEntry1_4 = tk.Entry(Frame2,textvariable=phon_no)
    TEntry1_4.place(x=449, y=485, height=34, width=342)
    TEntry1_4.configure(takefocus="",cursor="ibeam",font=('Helvetica', 20),relief='flat')

    email=tk.StringVar()
    TEntry1_5 = tk.Entry(Frame2,textvariable=email)
    TEntry1_5.place(x=449, y=559, height=35, width=342)
    TEntry1_5.configure(takefocus="",cursor="ibeam",font=('Helvetica', 20),relief='flat')

    address=tk.StringVar()
    TEntry1_6 = tk.Entry(Frame2,textvariable=address)
    TEntry1_6.place(x=449, y=650, height=35, width=342)
    TEntry1_6.configure(takefocus="",cursor="ibeam",font=('Helvetica', 20),relief='flat')

    city=tk.StringVar()
    TEntry1_6_1 = tk.Entry(Frame2,textvariable=city)
    TEntry1_6_1.place(x=449, y=742, height=35, width=342)
    TEntry1_6_1.configure(takefocus="",cursor="ibeam",font=('Helvetica', 20),relief='flat')

    state=tk.StringVar()
    TEntry1_6_2 = tk.Entry(Frame2,textvariable=state)
    TEntry1_6_2.place(x=449, y=828, height=35, width=342)
    TEntry1_6_2.configure(takefocus="",cursor="ibeam",font=('Helvetica', 20),relief='flat')


    def Submit_data():
        # user_info=['Name','password','age','gender','phon_no','email','Address','city','state']

        # ask_photo=tk.PhotoImage(file='')
        # data=[]
        # for field in range(0,len(user_info)):
        #     data.insert(field,input(user_info[field]+' ::'))
        if age.get()<12 or age.get()>120:                                               #check submitted data is valid or not
            tk.messagebox.showerror("Wrong Entry", "Enter age again")
            return None
        if phon_no.get()<6000000000 or phon_no.get()>10000000000:
            tk.messagebox.showerror("Wrong Entry", "Enter Phon Number again")
            return None

        data=[name.get(),password.get(),str(age.get()),gender.get(),str(phon_no.get()),email.get(),address.get(),city.get(),state.get()]
        
        # if data[3]==0:
        #     data[3]=='Male'
        # elif data[3]==1:
        #     data[3]=='Female'
        # else:
        #     data[3]=='Other'

        if data[0]=='' or data[1]=='' or data[6]=='' or data[7]=='' or data[8]=='':
            tk.messagebox.showerror("Blank field", "Fill All Form")
            return None
        print(data)
        with open('user_data.txt','a') as ud:
            # ud.write(str(data)+'\n')
            ud.write('\n')
            for i in data:
                ud.write(i+' ')
            
            New_home()

    submit = tk.Button(Frame2,command=Submit_data)
    submit.place(x=445, y=917, height=44, width=167)
    submit.configure(background="#ffffff",foreground="#000000",font=('Helvetica', 20),relief='flat',text='''Submit''')

    back= tk.Button(Frame2,command=lambda: New_home())
    back.place(x=234, y=917, height=44, width=167)
    back.configure(background="#ffffff",foreground="#000000",font=('Helvetica', 20),relief='flat',text='''Back''')
####################### New Driver ###################
def New_Driver(uemail,password):                                                    # For Job Proposol
    for widget in Frame2.winfo_children():
        widget.destroy()

    Label1 = tk.Label(Frame2)
    Label1.place(relx=0.0, rely=0.0, height=1080, width=1920)
    Label1.configure(background="#000000")
    Label1.configure(disabledforeground="#a3a3a3")
    Label1.configure(foreground="#ffffff")
    Label1.configure(image=Drive_register)
    Label1.configure(text='''Label''')

    name=tk.StringVar()
    Entry_Name = tk.Entry(Frame2,textvariable=name)
    Entry_Name.place(x=566,y=107,height=40,width=326)
    Entry_Name.configure( highlightthickness=0, borderwidth=0,background="black",foreground="white",font=('Helvetica', 20,'bold'))
    Entry_Name.configure(takefocus="",cursor="ibeam #ffffff",insertbackground='white')

    age=tk.IntVar()
    Spinbox1 = tk.Spinbox(Frame2, from_=1.0, to=120.0)
    Spinbox1.place(x=566,y=187,height=38,width=324)
    Spinbox1.configure(highlightthickness=0, borderwidth=0,background="black",foreground="white",relief='flat',font=('Helvetica', 20,'bold'))
    Spinbox1.configure(textvariable=age)

    style = ttk.Style()

    style.configure('TCombobox',foreground='white',releif='flat',background='black',activbackground='black',selectbackground="balck")
    style.configure('PSG.TCombobox', selectbackground='Black')
    def change_lable(object):
        w=TCombobox1.get()
        hide_by_lable.configure(text=w)
        

    gender=tk.StringVar()
    TCombobox1 = ttk.Combobox(Frame2,style="PSG.TCombobox")
    TCombobox1.place(x=566,y=265,height=39,width=324)
    TCombobox1.configure(textvariable=gender,takefocus="",font=('Helvetica', 20,'bold'))
    TCombobox1['values']=('Male','Female','Other')
    TCombobox1.current(0)
    TCombobox1.bind("<<ComboboxSelected>>",change_lable)

    w="Male"
    hide_by_lable=tk.Label(TCombobox1)
    hide_by_lable.place(x=0,y=0,height=39,width=307)
    hide_by_lable.configure(text=w,background='black',foreground='white',font=('Helvetica', 20,'bold'))

    hide1=tk.Label(Frame2)
    hide1.place(x=566,y=265,height=1,width=324)
    hide1.configure(background="black")

    hide2=tk.Label(Frame2)
    hide2.place(x=566,y=265,height=39,width=1)
    hide2.configure(background="black")

    hide3=tk.Label(Frame2)
    hide3.place(x=566,y=303,height=2,width=324)
    hide3.configure(background="black")

    hide4=tk.Label(Frame2)
    hide4.place(x=889,y=265,height=39,width=1)
    hide4.configure(background="black")

    phon_no=tk.IntVar()
    Entry_Phon_no = tk.Entry(Frame2,textvariable=phon_no)
    Entry_Phon_no.place(x=566,y=344,height=38,width=324)
    Entry_Phon_no.configure( highlightthickness=0, borderwidth=0,background="black",foreground="white",font=('Helvetica', 20,'bold'),takefocus="",cursor="ibeam #ffffff",insertbackground='white')

    email=tk.StringVar()
    Entry_email = tk.Entry(Frame2,textvariable=email)
    Entry_email.place(x=566,y=422,height=38,width=324)
    Entry_email.configure( highlightthickness=0, borderwidth=0,background="black",foreground="white",font=('Helvetica', 20,'bold'),takefocus="",cursor="ibeam #ffffff",insertbackground='white')


    language=tk.StringVar()
    Entry_address = tk.Entry(Frame2,textvariable=language)
    Entry_address.place(x=566,y=501,height=38,width=324)
    Entry_address.configure( highlightthickness=0, borderwidth=0,background="black",foreground="white",font=('Helvetica', 20,'bold'),takefocus="",cursor="ibeam #ffffff",insertbackground='white')

    address=tk.StringVar()
    Entry_language = tk.Entry(Frame2,textvariable=address)
    Entry_language.place(x=566,y=579,height=38,width=324)
    Entry_language.configure( highlightthickness=0, borderwidth=0,background="black",foreground="white",font=('Helvetica', 20,'bold'),takefocus="",cursor="ibeam #ffffff",insertbackground='white')

    lc_no=tk.StringVar()
    Entry_state = tk.Entry(Frame2,textvariable=lc_no)
    Entry_state.place(x=566,y=654,height=38,width=324)
    Entry_state.configure(highlightthickness=0, borderwidth=0,background="black",foreground="white",font=('Helvetica', 20,'bold'),takefocus="",cursor="ibeam #ffffff",insertbackground='white')

    # Entry_lc_date = ttk.Entry(Frame2,textvariable=lc_date)
    # Entry_lc_date.
    # Entry_lc_date.configure(takefocus="",cursor="ibeam")

    def show_date(object):
        w=str(lc_date.get())
        w=w.split('/')
        p=w[1]+'/'+w[0]+'/'+w[2]
        hide_by_lable1.configure(text=p,background='black',foreground='white',font=('Helvetica', 20,'bold'))

    # date1=tk.StringVar()
    lc_date=tk.StringVar()
    cal = DateEntry(Frame2,textvariable=lc_date,width=12, background='black',foreground='white', borderwidth=2, year=2021)
    cal.place(x=566,y=738,height=38,width=324)
    cal.bind("<<DateEntrySelected>>",show_date)
    # print(date1.get())
    today=date.today()
    w=str(today.day)+"/"+str(today.month)+"/"+str(today.year)
    hide_by_lable1=tk.Label(Frame2)
    hide_by_lable1.place(x=566,y=738,height=38,width=308)
    hide_by_lable1.configure(text=w,background='black',foreground='white',font=('Helvetica', 20,'bold'))


    def Submit_data():

        # user_info=['Name','password','age','gender','phon_no','email','Address','city','state']

        # ask_photo=tk.PhotoImage(file='')
        # data=[]
        # for field in range(0,len(user_info)):
        #     data.insert(field,input(user_info[field]+' ::'))
        if age.get()<12 or age.get()>120:
            tk.messagebox.showerror("Wrong Entry", "Enter age again")
            return None
        if phon_no.get()<6000000000 or phon_no.get()>10000000000:
            tk.messagebox.showerror("Wrong Entry", "Enter Phon Number again")
            return None

        data=[name.get(),str(age.get()),gender.get(),str(phon_no.get()),email.get(),address.get(),language.get(),lc_no.get(),lc_date.get()]
        
        if data[3]==0:
            data[3]=='Male'
        elif data[3]==1:
            data[3]=='Female'
        else:
            data[3]=='Other'

        if data[0]=='' or data[1]=='' or data[6]=='' or data[7]=='' or data[8]=='' :
            tk.messagebox.showerror("Blank field", "Fill All Form")
            return None
        print(data)
        with open('req_driver.txt','a') as ud:
            # ud.write(str(data)+'\n')
            for i in data:
                ud.write(i+',')
            ud.write('\n')
            New_After_login(uemail,password)

    Button1 = tk.Button(Frame2,command=Submit_data)
    Button1.place(x=558,y=830,height=43,width=275)
    Button1.configure(text="Submit",background='black',foreground='white',font=('Helvetica', 20,'bold'),relief='flat')

    back= tk.Button(Frame2,command=lambda: New_After_login(uemail,password))
    back.place(x=236, y=830, height=43, width=275)
    back.configure(background="Black",foreground="white",relief='flat',text='''Back''',font=('Helvetica', 20,'bold'))
################### Admin log in #####################
def admin_log_in():                                                                         #Admin menu
    for widget in Frame2.winfo_children():
            widget.destroy()
    Label1 = tk.Label(Frame2)
    Label1.place(relx=0.0, rely=0.0, height=1080, width=1920)
    Label1.configure(background="#6f6f64")
    Label1.configure(disabledforeground="#a3a3a3")
    Label1.configure(foreground="#ffffff")
    Label1.configure(image=admin_login)
    Label1.configure(text='''Label''')

    u_name=tk.StringVar()
    TEntry1 = tk.Entry(Frame2,textvariable=u_name)
    TEntry1.place(x=712,y=502, height=49,width=520)
    TEntry1.configure(takefocus="",cursor="ibeam",relief='flat',font=('Helvetica', 32,'bold'),background="white",foreground="black") 

    u_password=tk.StringVar()
    TEntry1_1 = tk.Entry(Frame2,textvariable=u_password,show="#")
    TEntry1_1.place(x=712,y=618, height=49,width=520)
    TEntry1_1.configure(takefocus="",cursor="ibeam",relief='flat',font=('Helvetica', 32,'bold'),background="white",foreground="black")
    def Login_admin():
        p=u_name.get()
        admin=tuple(['maulik','yash','yagnik'])
        password=['trip']
        if p==admin[0] or p==admin[1] or p==admin[2]:
            if u_password.get() == password[0]:
                # print("You Are In")
                Admin_menu()
                return None
        else:
            print('not')
    TButton1 = tk.Button(Frame2, command=Login_admin)
    TButton1.place(x=984,y=695, height=65,width=195)
    TButton1.configure(takefocus="",text='''SUBMIT''',font=('Copperplate Gothic Bold', 32,'bold'),background="white",relief='flat',foreground="#7ea8ab")

    back = tk.Button(Frame2, command=lambda: None)
    back.place(x=750,y=695, height=65, width=195)
    back.configure(takefocus="",text='''BACK''',font=('Copperplate Gothic Bold', 32,'bold'),background="White",relief='flat',foreground="#7ea8ab")
#################### Admin Menu ######################
def Admin_menu(): 
    for widget in Frame2.winfo_children():
            widget.destroy() 
    Label1 = tk.Label(Frame2)
    Label1.place(x=0.0, y=0.0, height=1015, width=1920)
    Label1.configure(background="#6f6f64")
    Label1.configure(disabledforeground="#a3a3a3")
    Label1.configure(foreground="#ffffff")
    Label1.configure(image=Admin_menu_ )
    Label1.configure(text='''Label''')

    Drive_list = tk.Button(Frame2,command=lambda: New_display("driver",16))
    Drive_list.place(x=1500, y=85,height=63, width=320)
    Drive_list.configure(activebackground="#544d47",activeforeground="#ddc3a5",font=('Helvetica', 32,'bold'),background="#544d47",disabledforeground="#a3a3a3",foreground="#ddc3a5",highlightbackground="#544d47",highlightcolor="black",relief='flat',text='''Driver List''')

    New_drive = tk.Button(Frame2,command=lambda: edit_data('places',20))#creat_new_drive())
    New_drive.place(x=1500, y=230,height=63, width=320)
    New_drive.configure(activebackground="#544d47",activeforeground="#ddc3a5",font=('Helvetica', 31,'bold'),background="#544d47",disabledforeground="#a3a3a3",foreground="#ddc3a5",highlightbackground="#544d47",highlightcolor="black",relief='flat',text='''Add New Places''')

    Edit_driver = tk.Button(Frame2,command=lambda: Add_New_driver())#edit_data("driver",10))
    Edit_driver.place(x=1500, y=375,height=63, width=320)
    Edit_driver.configure(activebackground="#544d47",activeforeground="#ddc3a5",font=('Helvetica', 32,'bold'),background="#544d47",disabledforeground="#a3a3a3",foreground="#ddc3a5",highlightbackground="#544d47",highlightcolor="black",relief='flat',text='''Edit Driver''')

    Custmor_list = tk.Button(Frame2,command=lambda: New_display("user_data",22))#Display("user_data",15))
    Custmor_list.place(x=1500, y=520,height=63, width=320)
    Custmor_list.configure(activebackground="#544d47",activeforeground="#ddc3a5",font=('Helvetica', 32,'bold'),background="#544d47",disabledforeground="#a3a3a3",foreground="#ddc3a5",highlightbackground="#544d47",highlightcolor="black",relief='flat',text='''Customer List''')

    edit_custmor = tk.Button(Frame2,command=lambda: edit_data("user_data",26))#edit_data('use_info',12))
    edit_custmor.place(x=1500, y=665,height=63, width=320)
    edit_custmor.configure(activebackground="#544d47",activeforeground="#ddc3a5",font=('Helvetica', 32,'bold'),background="#544d47",disabledforeground="#a3a3a3",foreground="#ddc3a5",highlightbackground="#544d47",highlightcolor="black",relief='flat',text='''Edit Customer''')

    All_trip = tk.Button(Frame2,command=lambda: New_display("history",18))#Display("history",11))
    All_trip.place(x=1500, y=796,height=63, width=320)
    All_trip.configure(activebackground="#544d47",activeforeground="#ddc3a5",font=('Helvetica', 32,'bold'),background="#544d47",disabledforeground="#a3a3a3",foreground="#ddc3a5",highlightbackground="#544d47",highlightcolor="black",relief='flat',text='''All Trips''')

    back = tk.Button(Frame2,command=lambda: New_home())
    back.place(x=1500, y=896,height=63, width=320)
    back.configure(activebackground="white",activeforeground="black",font=('Helvetica', 32,'bold'),background="white",disabledforeground="white",foreground="black",highlightbackground="white",highlightcolor="black",relief='flat',text='''Back''')
###################  Display #########################
def New_display(filename,font1):                                
    for widget in Frame2.winfo_children():                      #displays Requried Data
            widget.destroy()

    Label1 = tk.Label(Frame2)
    Label1.place(relx=0.0, rely=0.0, height=1030, width=1920)
    Label1.configure(background="#6f6f64")
    Label1.configure(disabledforeground="#a3a3a3")
    Label1.configure(foreground="#ffffff")
    Label1.configure(text='''Label''')
    if filename=='driver':
        Label1.configure(image=Display_driver)
    elif filename=="user_data":
        Label1.configure(image=Display_Users)
    else:
        Label1.configure(image=Display_history)
        
    
    Text1 = tk.Text(Frame2)
    Text1.place(x=150, y=170, height=760, width=1650)
    Text1.configure(background="white",font=('Consolas',font1,'bold'),foreground="black",highlightbackground="#d9d9d9",highlightcolor="black",insertbackground="black",selectbackground="blue",selectforeground="white",wrap="word")

    Button1 = tk.Button(Frame2,command=lambda: Admin_menu())#after_Admin_login())
    Button1.place(x=810, y=960, height=40, width=300)
    Button1.configure(activebackground="#ececec",activeforeground="#000000",background="white",disabledforeground="#a3a3a3",font="-family {Segoe UI} -size 25",foreground="#000000",highlightbackground="#d9d9d9",highlightcolor="black",relief='flat',text='''Back''')

    # Label3 = tk.Label(Frame2)
    # Label3.place(x=66, y=80, height=21, width=794)
    # Label3.configure(background="#d9d9d9",disabledforeground="#a3a3a3",font="-family {Constantia} -size 8 -weight bold",foreground="#000000")
    # Label3.configure(text='''Refe No   Name   source    Destination            Date          &      Time   Type           Price    Total Distance    Driver Name          Texi Number''')

    with open(f"{filename}.txt",'r',encoding='UTF-8') as hi:
        for i in hi.readlines():
            if len(i) > 1 or i != '\n':
                # print(i)
                if filename=="driver":
                    # print(len(i))
                    i=i.split(' ')
                    if i[9]=='GUJ':
                        i.insert(9,"----")
                    p="{:12s}{:4s}{:6s}{:14s}{:25s}{:15s}{:8s}{:8s}{:8s}{:8s}{:6s}{:14s}{:8s}".format(i[0],i[1],i[2],i[3],i[4],i[5],i[6],i[7],i[8],i[9],i[10],i[11],i[12])
                elif filename=="user_data":
                    i=i.split(' ')
                    if i[0]!=' ' and i[0]!='\n' and len(i[0])>3:
                        p="{:10s}{:10s}{:6s}{:8s}{:15s}{:20s}{:13s}{:10s}{:8s}".format(i[0],i[1],i[2],i[3],i[4],i[5],i[6],i[7],i[8])
                else :
                    i=i.split(',')
                    p="{:8s}{:10s}{:12s}{:12s}{:17s}{:18s}{:11s}{:10s}{:14s}{:12s}\n".format(i[0],i[1],i[2],i[3],i[4],i[5],i[6],i[7],i[8],i[9])
                    
                Text1.insert(tk.END,p)
                Text1.insert(tk.END,"\n")
        Text1.configure(state=tk.DISABLED)
##################  Edit Driver ######################
def Add_New_driver():                                            
    for widget in Frame2.winfo_children():
        widget.destroy()

    Label1 = tk.Label(Frame2)
    Label1.place(relx=0.0, rely=0.0, height=1030, width=1920)
    Label1.configure(background="#6f6f64")
    Label1.configure(disabledforeground="#a3a3a3")
    Label1.configure(foreground="#ffffff")
    Label1.configure(image=Add_driver)
    Label1.configure(text='''Label''')


    Text1 = tk.Text(Frame2)
    Text1.place(x=147, y=165, height=367, width=1650)
    Text1.configure(background="white",font=('Consolas',16,'bold'),foreground="black",highlightbackground="#d9d9d9",highlightcolor="black",insertbackground="black",selectbackground="blue",selectforeground="white",wrap="word")

    Button1 = tk.Button(Frame2,command=lambda: save())
    Button1.place(x=810, y=960, height=40, width=300)
    Button1.configure(activebackground="#ececec",activeforeground="#000000",background="white",disabledforeground="#a3a3a3",font="-family {Segoe UI} -size 25",foreground="#000000",highlightbackground="#d9d9d9",highlightcolor="black",relief='flat',text='''Save And Back''')



    with open(f"req_driver.txt",'r',encoding='UTF-8') as hi:
        data=hi.read()
        data=data.split("\n")
        for i in data:
            i=i.split(',')
            
            for p in i:
                p=p+' '
                Text1.insert(tk.END,p)
            Text1.insert(tk.END,"\n")


    Text2 = tk.Text(Frame2)
    Text2.place(x=147, y=580, height=370, width=1650)
    Text2.configure(background="white",font=('Consolas',16,'bold'),foreground="black",highlightbackground="#d9d9d9",highlightcolor="black",insertbackground="black",selectbackground="blue",selectforeground="white",wrap="word") #Maulik_GUI

    with open(f"driver.txt",'r',encoding='UTF-8') as hi:
        data=hi.read()
        data=data.split("\n")
        for i in data:
            i=i.split(' ')
            for p in i:
                p=p+' '
                Text2.insert(tk.END,p)
            Text2.insert(tk.END,"\n")

    def save():
            with open(f"driver.txt","w") as uf:
                uf.write(Text2.get('1.0',tk.END))
                Admin_menu()    
##################  Edit Data   ######################
def edit_data(filename,font2):
    for widget in Frame2.winfo_children():
        widget.destroy()
        #Maulik_GUI
    Label1 = tk.Label(Frame2)
    Label1.place(relx=0.0, rely=0.0, height=1030, width=1920)
    Label1.configure(background="#000000")
    Label1.configure(disabledforeground="#a3a3a3")
    Label1.configure(foreground="#ffffff")
    Label1.configure(image=edit)
    Label1.configure(text='''Label''')
    
    TSeparator1 = ttk.Separator(Frame2)
    TSeparator1.place(x=42, y=72, width=0)

    Text1 = tk.Text(Frame2)
    Text1.place(x=150, y=170, height=760, width=1650)
    Text1.configure(background="white",font=('Consolas',font2,'bold'),foreground="black",highlightbackground="#d9d9d9",highlightcolor="black",insertbackground="black",selectbackground="blue",selectforeground="white",wrap="word")

    Button1 = tk.Button(Frame2,command=lambda:save())
    Button1.place(x=810, y=960, height=40, width=300)
    Button1.configure(activebackground="#ececec",activeforeground="#000000",background="white",disabledforeground="#a3a3a3",font="-family {Segoe UI} -size 25",foreground="#000000",highlightbackground="#d9d9d9",highlightcolor="black",relief='flat',text='''Save And Back''')

    # Label3 = tk.Label(Frame2)
    # Label3.place(x=66, y=80, height=21, width=794)
    # Label3.configure(background="#d9d9d9",disabledforeground="#a3a3a3",font="-family {Constantia} -size 10 -weight bold",foreground="#000000")
    # Label3.configure(text='''Refe No   Name   source  Destination           Date         &      Time   Type           Price    Total Distance    Driver Name          Texi Number''')

    with open(f"{filename}.txt",'r',encoding='UTF-8') as hi:
        data=hi.read()
        data=data.split("\n")
        for i in data:
            i=i.split(',')
            for p in i:
                p=p+' '
                Text1.insert(tk.END,p)
            Text1.insert(tk.END,"\n")
            
    def save():
        with open(f"{filename}.txt","w") as uf:
            data=Text1.get('1.0',tk.END).splitlines()
            for line in data:
                if len(line)<5 or line=='\n':
                    continue
                uf.write(line[0:-1])
                uf.write('\n')
            Admin_menu()    
##################  History user #####################
def History_user(uemail,password):         
    uname=''
    for widget in Frame2.winfo_children():
        widget.destroy()

    Label1 = tk.Label(Frame2)
    Label1.place(relx=0.0, rely=0.0, height=1030, width=1920)
    Label1.configure(background="#6f6f64")
    Label1.configure(disabledforeground="#a3a3a3")
    Label1.configure(foreground="#ffffff")

    with open("user_data.txt",'r') as ud:
        for i in ud.readlines():
            if len(i) > 1 or i != '\n':
                    i=i.split(' ')
                    if uemail==i[5]:
                        uname=i[0]
                        break 
    Label1.configure(image=Display_history)
    Label1.configure(text='''Label''')

    Text1 = tk.Text(Frame2)
    Text1.place(x=150, y=170, height=760, width=1645)
    Text1.configure(background="white",font=('Consolas',18,'bold'),foreground="black",highlightbackground="#d9d9d9",highlightcolor="black",insertbackground="black",selectbackground="blue",selectforeground="white",wrap="word")

    Button1 = tk.Button(Frame2,command=lambda: New_After_login(uemail,password))
    Button1.place(x=810, y=960, height=40, width=300)
    Button1.configure(activebackground="#ececec",activeforeground="#000000",background="white",disabledforeground="#a3a3a3",font="-family {Segoe UI} -size 25",foreground="#000000",highlightbackground="#d9d9d9",highlightcolor="black",relief='flat',text='''Back''')

    with open("history.txt",'r',encoding='UTF-8') as hi:
        data=hi.read()
        data=data.split("\n")
        for i in data:
            if len(i)>10:
                i=i.split(',')
                if uname == i[1] and len(i[0])>2:
                    p="{:8s}{:8s}{:10s}{:13s}{:20s}{:12s}{:11s}{:10s}{:14s}{:12s}".format(i[0],i[1],i[2],i[3],i[4],i[5],i[6],i[7],i[8],i[9])
                    # print(p)
                    Text1.insert(tk.END,p)
                    Text1.insert(tk.END,"\n")
            
        Text1.configure(state=tk.DISABLED)

def __main__():
    New_home()

__main__()
win.mainloop()